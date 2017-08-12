# Author: Colton Sweeney
# Date: August 10, 2017
# Name: TrustRadius Review Scraper
#
# Lang: Python
# Dependencies: BeautifulSoup4
#               Requests
#
# Description: As part of my current workload, I have to create
#              weekly reports gathered from reviews on TrustRadius.
#              In order to expedite this process, I decided to create 
#              simple Python bot that scours the TrustRadius page for
#              all new reviews and then creates detailed reports directly 
#              from the information gathered in those reports.

# import required dependencies
from bs4 import BeautifulSoup
import requests
import docx

# Open up TrustRadius page for Sitefinity
url = "https://www.trustradius.com/products/progress-sitefinity/reviews"
data = requests.get(url).text
soup = BeautifulSoup(data, "html.parser")

# Find all review links from main product page
query = soup.find_all('a', {
    'class' : 'link-to-review-btn'
})

# List that holds URLs of individual reviews
links = []

# Append all the URLs to the list we created
for link in query:
    links.append('https://www.trustradius.com' + link.get('href'))

# Print number of reviews to command line
print "%d links found.\n" % len(links)

# Create custom data structure Review to hold user reviews
class Review:
    def __init__(self, name, position, company, rating, goodies):
        self.name = name
        self.position = position
        self.company = company
        self.rating = rating
        self.goodies = goodies


# Function for returning review sections from review page
# return (dictionary): a key-value list of the headings and review text
# parameters: link (string): the url for the review
def findMaterials(link):
    # Parse the given link into some Beautiful Soup
    req = requests.get(link).text
    reviews = BeautifulSoup(req, 'html.parser')

    # Set up list string variables.
    reviewAuthor = []
    reviewPosition = []
    reviewCompany = []
    reviewRating = []
    sectionHeading = []
    sectionText = []

    # Find the authors name (if there is one)
    for review in reviews.find_all('span', {'itemprop': 'author'}):
        reviewAuthor.append(review.contents[0].text)
    
    # Find the author's position and company (if applicable)
    for review in reviews.find_all('span', {'class': 'user-info'}):
        reviewPosition.append(review.contents[0].text)
        reviewCompany.append(review.contents[1].text)

    # Find what the user rated Sitefinity
    reviewRating = reviews.find_all('span', class_='number')[0].text

    # Perform find.contents[] for all of the headings and text
    # and append them to our functions variables
    for review in reviews.find_all('div', {'class': 'description'}):
        # Recieve review section headings
        sectionHeading.append(review.contents[0].contents[0].contents[1].contents[0].contents[0].contents[0].contents[0].contents[0].text)
        sectionHeading.append(review.contents[0].contents[0].contents[1].contents[1].contents[0].contents[0].contents[0].contents[0].text)
        sectionHeading.append(review.contents[0].contents[0].contents[1].contents[2].contents[0].contents[0].contents[0].contents[0].text)
        sectionHeading.append(review.contents[0].contents[0].contents[1].contents[3].contents[0].contents[0].contents[0].contents[0].text)        
        sectionHeading.append(review.contents[0].contents[0].contents[1].contents[4].contents[0].contents[0].contents[0].contents[0].text)
        sectionHeading.append(review.contents[0].contents[0].contents[1].contents[5].contents[0].contents[0].contents[0].contents[0].text)

        # Recieve review text
        sectionText.append(review.contents[0].contents[0].contents[1].contents[0].contents[1].contents[0].contents[0].contents[0].text)
        sectionText.append(review.contents[0].contents[0].contents[1].contents[1].contents[1].contents[0].contents[0].contents[0].text)
        sectionText.append(review.contents[0].contents[0].contents[1].contents[2].contents[1].contents[0].contents[0].contents[0].text)
        sectionText.append(review.contents[0].contents[0].contents[1].contents[3].contents[1].contents[0].contents[0].contents[0].text)
        sectionText.append(review.contents[0].contents[0].contents[1].contents[4].contents[1].contents[0].contents[0].contents[0].text)
        sectionText.append(review.contents[0].contents[0].contents[1].contents[5].contents[1].contents[0].contents[0].contents[0].text)

    # Wrap up the review information into a dictionary, this is for easy handling    
    reviewDict = dict(zip(sectionHeading, sectionText))

    # Create a new review using our Review class, and return that review
    rev = Review(reviewAuthor, reviewPosition, reviewCompany, reviewRating, reviewDict)
    return rev



reviewGuide = []
for num in range(len(links)):
    reviewGuide.append(findMaterials(links[num]))


doc = docx.Document()
doc.add_heading('Trust Radius Weekly Report', 0)

def createPage(page):
    doc.add_heading(page.name, 1)
    doc.add_heading(page.rating + ' out of 10', 1)

    for x, y in page.goodies.items():
        doc.add_heading(x, 2)
        doc.add_paragraph(y)

    doc.add_page_break()

for review in reviewGuide:
    createPage(review)

print('Successfully created a .docx with %d reviews. Check out results.docx...' % len(links))
doc.save('results.docx')


